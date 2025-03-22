import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from io import BytesIO, StringIO
import cv2
import numpy as np
from paddleocr import PaddleOCR
from PIL import Image
from modules.utils import merge_predictions
from groq import Groq
import torch
import re
import pytesseract
import difflib

SENSITIVITY_LEVELS = {
    "Low": {"I-ACCOUNTNUM", "I-CREDITCARDNUMBER", "I-DRIVERLICENSENUM", "I-IDCARDNUM", "I-PASSWORD", "I-SOCIALNUM", "I-TAXNUM"},
    "Medium": {"I-ACCOUNTNUM", "I-CREDITCARDNUMBER", "I-DRIVERLICENSENUM", "I-IDCARDNUM", "I-PASSWORD", "I-SOCIALNUM", "I-TAXNUM",
               "I-EMAIL", "I-GIVENNAME", "I-SURNAME", "I-TELEPHONENUM", "I-USERNAME"},
    "High": {"I-ACCOUNTNUM", "I-BUILDINGNUM", "I-CITY", "I-CREDITCARDNUMBER", "I-DATEOFBIRTH", "I-DRIVERLICENSENUM", "I-EMAIL",
             "I-GIVENNAME", "I-IDCARDNUM", "I-PASSWORD", "I-SOCIALNUM", "I-STREET", "I-SURNAME", "I-TAXNUM", "I-TELEPHONENUM",
             "I-USERNAME", "I-ZIPCODE"}
}

# Regex patterns for PII detection as a fallback
PII_PATTERNS = {
    "I-EMAIL": r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+",
    "I-TELEPHONENUM": r"\b(\+\d{1,3}[- ]?)?(\(?\d{3}\)?[- ]?)?\d{3}[- ]?\d{4}\b",
    "I-SOCIALNUM": r"\b\d{3}-\d{2}-\d{4}\b",
    "I-ZIPCODE": r"\b\d{5}(-\d{4})?\b",
    "I-CREDITCARDNUMBER": r"\b\d{4}-\d{4}-\d{4}-\d{4}\b",
}

def regex_redact(text, allowed_labels):
    """Fallback redaction using regex for PII types."""
    redacted_text = text
    for label, pattern in PII_PATTERNS.items():
        if label in allowed_labels:
            redacted_text = re.sub(pattern, "[REDACTED]", redacted_text)
    return redacted_text

def groq_redact_remaining(text, client, model_name, sensitivity, allowed_labels):
    if not client or not model_name:
        return text
    # Enhanced prompt with examples for better redaction
    prompt = f"""
    You are an expert in data privacy and entity recognition. Your task is to identify and redact sensitive information in the input text based on the specified PII types for the given sensitivity level. Replace identified entities with '[REDACTED]'. Preserve original formatting, whitespace, and punctuation. Do not redact entities already marked with '[REDACTED]'.

    **Sensitivity Level**: {sensitivity}
    **PII Types to Redact**: {', '.join(sorted(allowed_labels))}
    **Examples of PII to Redact**:
    - I-GIVENNAME: John, Jane
    - I-SURNAME: Doe, Smith
    - I-EMAIL: john.doe@example.com
    - I-TELEPHONENUM: (555) 123-4567, +1 555-123-4567
    - I-STREET: 123 Main Street
    - I-CITY: Anytown
    - I-ZIPCODE: 91234
    - I-SOCIALNUM: 123-45-6789
    **Input Text**:
    {text}

    **Output Format**:
    - Return only the redacted text, no additional commentary or labels. No "here's the redacted text" sentence too, only the redacted text. 
    """
    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "You are a precise and efficient assistant specializing in text redaction."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1024,
            temperature=0.1
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        raise Exception(f"Groq API error: {str(e)}")

def redact_text(text, predictions, allowed_labels, client=None, model_name=None):
    # Step 1: NER-based redaction
    merged_predictions = merge_predictions(predictions)
    merged_predictions = sorted(merged_predictions, key=lambda x: x['start'], reverse=True)
    redacted_text = list(text)
    for pred in merged_predictions:
        if pred['entity'] in allowed_labels:
            start, end = pred['start'], pred['end']
            for i in range(start, end):
                redacted_text[i] = ''
            redacted_text[start] = '[REDACTED]'
    ner_redacted_text = ''.join(redacted_text)

    # Step 2: Groq API refinement
    if client and model_name:
        try:
            groq_redacted_text = groq_redact_remaining(
                ner_redacted_text, client, model_name,
                "Custom" if allowed_labels != set(SENSITIVITY_LEVELS.get("Medium", set())) else "Medium",
                allowed_labels
            )
        except Exception as e:
            print(f"Groq API failed: {str(e)}. Falling back to regex.")
            groq_redacted_text = ner_redacted_text
    else:
        groq_redacted_text = ner_redacted_text

    # Step 3: Regex fallback for missed PII
    final_redacted_text = regex_redact(groq_redacted_text, allowed_labels)
    return final_redacted_text

def redact_pdf(pipe, file, sensitivity, custom_pii_types=None, client=None, model_name=None):
    doc = fitz.open(stream=file.read(), filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text() + "\n"
    
    redacted_text = redact_text(full_text, pipe(full_text), 
                               SENSITIVITY_LEVELS.get(sensitivity, SENSITIVITY_LEVELS["Medium"]) if sensitivity != "Custom" else set(custom_pii_types),
                               client, model_name)
    
    new_doc = fitz.open()
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
        new_page.insert_text((50, 50), redacted_text, fontsize=12)
    
    output = BytesIO()
    new_doc.save(output)
    output.seek(0)
    return output, redacted_text

def redact_txt(pipe, file, sensitivity, custom_pii_types=None, client=None, model_name=None):
    text = file.read().decode('utf-8')
    return redact_text(text, pipe(text),
                       SENSITIVITY_LEVELS.get(sensitivity, SENSITIVITY_LEVELS["Medium"]) if sensitivity != "Custom" else set(custom_pii_types),
                       client, model_name)

def redact_xlsx(pipe, file, sensitivity, custom_pii_types=None, client=None, model_name=None):
    df = pd.read_excel(file)
    allowed_labels = SENSITIVITY_LEVELS.get(sensitivity, SENSITIVITY_LEVELS["Medium"]) if sensitivity != "Custom" else set(custom_pii_types)
    for col in df.columns:
        for idx, cell in enumerate(df[col]):
            if pd.notna(cell):
                cell_text = str(cell)
                redacted_cell = redact_text(cell_text, pipe(cell_text), allowed_labels, client, model_name)
                df.at[idx, col] = redacted_cell
    output = StringIO()
    df.to_csv(output, index=False)
    output.seek(0)
    return output

def redact_docx(pipe, file, sensitivity, custom_pii_types=None, client=None, model_name=None):
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    redacted_text = redact_text(text, pipe(text),
                               SENSITIVITY_LEVELS.get(sensitivity, SENSITIVITY_LEVELS["Medium"]) if sensitivity != "Custom" else set(custom_pii_types),
                               client, model_name)
    
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = redacted_text if '[REDACTED]' in redacted_text else run.text
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output, redacted_text

def redact_image(pipe, file, sensitivity="Medium", custom_pii_types=None, client=None, model_name=None):

    """
    https://medium.com/geekculture/tesseract-ocr-understanding-the-contents-of-documents-beyond-their-text-a98704b7c655
    """

    # Set Tesseract path (adjust based on your system)
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    
    # Load and preprocess image
    img = Image.open(file)
    img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    # Enhanced preprocessing: Denoise, contrast adjustment, and adaptive thresholding
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    gray = cv2.convertScaleAbs(gray, alpha=2.0, beta=0)  # Increased contrast further
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 21, 5)
    print("Preprocessing completed: Image converted to thresholded grayscale.")  # Debug: Preprocessing step
    
    # Extract text with OCR
    text = pytesseract.image_to_string(thresh, config='--psm 6 --oem 3')
    print("Extracted Text:", text)  # Debug: Print extracted text
    
    # Detect PII in the extracted text
    predictions = pipe(text)
    merged_predictions = merge_predictions(predictions)
    print("Raw Predictions:", predictions)  # Debug: Print raw PII predictions
    print("Merged Predictions:", merged_predictions)  # Debug: Print merged PII predictions
    
    # Determine which PII types to redact
    allowed_labels = SENSITIVITY_LEVELS.get(sensitivity, SENSITIVITY_LEVELS["Medium"])
    if sensitivity == "Custom" and custom_pii_types:
        allowed_labels = set(custom_pii_types)
    print("Allowed Labels:", allowed_labels)  # Debug: Print allowed PII types
    
    # Get bounding box data for all text in the image
    boxes = pytesseract.image_to_data(thresh, output_type=pytesseract.Output.DICT, config='--psm 6 --oem 3')
    print("Bounding Boxes:", boxes)  # Debug: Print bounding box data
    
    # Apply redaction to PII regions
    redacted_regions = []
    for pred in merged_predictions:
        if pred['entity'] in allowed_labels:
            pred_text = text[pred['start']:pred['end']].lower().strip()
            print(f"Processing PII: {pred_text} (Entity: {pred['entity']}, Start: {pred['start']}, End: {pred['end']})")  # Enhanced debug
            # Find matching bounding boxes for the predicted PII
            matched_indices = []
            for i, word in enumerate(boxes['text']):
                if boxes['conf'][i] > 40 and word:  # Further lowered confidence threshold
                    word_lower = word.lower().strip()
                    similarity = difflib.SequenceMatcher(None, word_lower, pred_text).ratio()
                    # Broadened matching criteria with substring and partial matches
                    if (similarity > 0.4 or 
                        word_lower in pred_text or 
                        pred_text in word_lower or 
                        any(pred_text.startswith(w) or pred_text.endswith(w) for w in word_lower.split())):
                        x, y, w, h = boxes['left'][i], boxes['top'][i], boxes['width'][i], boxes['height'][i]
                        print(f"Redacting: {word} at ({x}, {y}, {w}, {h}) with similarity {similarity}")  # Debug: Log redaction
                        redacted_regions.append((x, y, w, h))
                        matched_indices.append(i)
            if not matched_indices:
                print(f"Warning: No matching bounding box found for PII: {pred_text} (Entity: {pred['entity']})")
    
    # Handle multi-word PII by merging overlapping regions with relaxed threshold
    if redacted_regions:
        redacted_regions.sort(key=lambda r: (r[1], r[0]))  # Sort by y then x for better alignment
        merged_regions = []
        current_region = list(redacted_regions[0])
        for next_region in redacted_regions[1:]:
            # Relaxed overlap condition for multi-word entities
            if (current_region[1] + current_region[3] >= next_region[1] - 20 and 
                abs(current_region[0] - next_region[0]) < 200):  # Increased tolerance
                current_region[2] = max(current_region[0] + current_region[2], next_region[0] + next_region[2]) - current_region[0]
                current_region[3] = max(current_region[1] + current_region[3], next_region[1] + next_region[3]) - current_region[1]
            else:
                merged_regions.append(tuple(current_region))
                current_region = list(next_region)
        merged_regions.append(tuple(current_region))
        
        # Reapply blurring to merged regions with padding
        for x, y, w, h in merged_regions:
            # Add padding to ensure full coverage
            padding = 10
            y_start = max(0, y - padding)
            y_end = min(img_cv.shape[0], y + h + padding)
            x_start = max(0, x - padding)
            x_end = min(img_cv.shape[1], x + w + padding)
            roi = img_cv[y_start:y_end, x_start:x_end]
            blurred = cv2.GaussianBlur(roi, (111, 111), 0)  # Increased blur kernel
            img_cv[y_start:y_end, x_start:x_end] = blurred
            print(f"Merged Redaction at ({x_start}, {y_start}, {x_end - x_start}, {y_end - y_start}) with padding")  # Debug: Log merged regions

    # Save the redacted image
    output = BytesIO()
    Image.fromarray(cv2.cvtColor(img_cv, cv2.COLOR_BGR2RGB)).save(output, format="PNG")
    output.seek(0)
    return output